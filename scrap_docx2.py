import requests
from bs4 import BeautifulSoup
import docx2txt
import docx

# URL halaman website yang ingin disalin
url = 'https://www.dicoding.com/academies/261/tutorials/14107'

# Lakukan permintaan HTTP ke URL
response = requests.get(url)

# Buat objek BeautifulSoup dari respons HTTP
soup = BeautifulSoup(response.content, 'html.parser')

# Dapatkan semua teks dari halaman website
text = soup.get_text()

# Dapatkan semua video dari halaman website
videos = soup.findAll('iframe', {'class': 'embed-responsive-item'})

# Buat dokumen docx
doc = docx.Document()

# Tambahkan judul dan sub judul untuk setiap bagian teks
title_element = soup.find('h1')
sub_title_element = soup.find('h2')

title = title_element.text if title_element else 'Tidak ada judul'
sub_title = sub_title_element.text if sub_title_element else 'Tidak ada sub judul'

doc.add_paragraph('Judul: ' + title)
doc.add_paragraph('Sub Judul: ' + sub_title)


# Tambahkan teks ke dokumen docx
for paragraph in text.split('\n'):
    doc.add_paragraph(paragraph)

# Tambahkan video ke dokumen docx
for video in videos:
    image_url = video['src'].split('?')[0]
    link_url = video['src']

    doc.add_paragraph('Video: ' + video['src'])
    doc.add_picture(image_url)
    doc.add_paragraph('Link: ' + link_url)

# Simpan dokumen docx
doc.save('output.docx')
